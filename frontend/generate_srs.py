"""Generate SRS document as .docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    s = doc.styles[f'Heading {level}']
    s.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def add_table(headers, rows_data):
    table = doc.add_table(rows=len(rows_data)+1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows_data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    return table

# ── Title Page ──
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Software Requirements Specification'); r.bold = True; r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI Health Insurance Cold Calling Agent'); r.font.size = Pt(18)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 2.0.0'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared by: NCAI Engineering Team\nDate: June 2026'); r.font.size = Pt(12)

doc.add_page_break()

# ── TOC ──
doc.add_heading('Table of Contents', level=1)
toc = ['1. Introduction','   1.1 Purpose','   1.2 Document Conventions','   1.3 Intended Audience',
       '   1.4 Scope','   1.5 Definitions & Acronyms','2. Overall Description','   2.1 Product Perspective',
       '   2.2 Product Functions','   2.3 User Characteristics','   2.4 Assumptions & Dependencies',
       '3. System Features & Requirements','   3.1 Authentication & Authorization','   3.2 Lead Management',
       '   3.3 Campaign Management','   3.4 Call Management','   3.5 AI Conversation Engine',
       '   3.6 Text-to-Speech (TTS)','   3.7 Speech-to-Text (STT)','   3.8 Real-Time Dashboard',
       '   3.9 RAG Knowledge Base','   3.10 Analytics & Reporting','4. External Interface Requirements',
       '   4.1 User Interfaces','   4.2 Hardware Interfaces','   4.3 Software Interfaces',
       '   4.4 Communication Interfaces','5. Non-Functional Requirements','   5.1 Performance','   5.2 Security',
       '   5.3 Reliability','   5.4 Scalability','   5.5 Maintainability','6. System Architecture',
       '7. API Documentation','8. Database Design','9. Deployment']
for t in toc:
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Introduction ──
doc.add_heading('1. Introduction', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph('The AI Health Insurance Cold Calling Agent is an intelligent outbound calling system designed to automate health insurance lead qualification through natural AI-powered conversations. It replaces manual cold calling with an automated pipeline that calls leads, engages them in natural dialogue, qualifies their interest, and logs results in real-time.')

doc.add_heading('1.2 Document Conventions', level=2)
for b in ['Bold text indicates key terms and important concepts', 'Monospace text indicates code, API endpoints, and file paths', 'REQ-XXX format identifies unique requirement identifiers']:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('1.3 Intended Audience', level=2)
for a in ['Development team implementing the system', 'QA engineers writing test cases', 'Project managers tracking requirements', 'Stakeholders evaluating system capabilities', 'DevOps engineers managing deployment']:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('1.4 Scope', level=2)
doc.add_paragraph('The system encompasses a complete AI-driven cold calling solution including: AI conversation engine with multi-provider LLM support (Groq, Gemini, Ollama), text-to-speech synthesis (ElevenLabs, gTTS), speech-to-text transcription (Whisper), telephony integration (SignalWire/Twilio), real-time WebSocket dashboard, lead and campaign management, JWT-based authentication, RAG knowledge base, and comprehensive analytics. The system is deployed via Docker Compose and Render.')

doc.add_heading('1.5 Definitions & Acronyms', level=2)
add_table(['Term', 'Definition'], [
    ('LLM', 'Large Language Model - AI model for natural language generation'),
    ('TTS', 'Text-to-Speech - converts text to spoken audio'),
    ('STT', 'Speech-to-Text - converts audio to text transcription'),
    ('RAG', 'Retrieval-Augmented Generation - enhances LLM with document context'),
    ('JWT', 'JSON Web Token - stateless authentication mechanism'),
    ('TwiML', 'Twilio Markup Language - XML for telephony call flow'),
    ('SIP', 'Session Initiation Protocol - VoIP signaling protocol'),
    ('PBX', 'Private Branch Exchange - internal telephone network'),
    ('ngrok', 'Public HTTP tunnel for local webhook endpoints'),
    ('Groq', 'Cloud LLM inference provider - free tier, high speed'),
])
doc.add_page_break()

# ── 2. Overall Description ──
doc.add_heading('2. Overall Description', level=1)
doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph('The system is a standalone web application with a React frontend and FastAPI backend. It interfaces with external services (SignalWire for telephony, Groq/Gemini for LLM, ElevenLabs for TTS) but is designed to fall back to local alternatives (gTTS, Whisper, Ollama) when external services are unavailable. The system stores all data in a PostgreSQL database and provides real-time updates via WebSocket connections.')

doc.add_heading('2.2 Product Functions', level=2)
for f in ['REQ-FUNC-001: User authentication via JWT tokens with automatic refresh',
'REQ-FUNC-002: Lead CRUD operations with CSV bulk upload',
'REQ-FUNC-003: Campaign creation, activation, and tracking',
'REQ-FUNC-004: AI-powered outbound calling with natural conversation',
'REQ-FUNC-005: Multi-provider LLM with automatic fallback (Groq, Gemini, Ollama, OpenAI, Anthropic)',
'REQ-FUNC-006: Text-to-speech with ElevenLabs and gTTS fallback',
'REQ-FUNC-007: Speech-to-text via local Whisper model',
'REQ-FUNC-008: Real-time dashboard with WebSocket updates',
'REQ-FUNC-009: Call history with full transcript and sentiment analysis',
'REQ-FUNC-010: RAG knowledge base for insurance document context',
'REQ-FUNC-011: Lead qualification scoring and status tracking',
'REQ-FUNC-012: Local test endpoints for AI pipeline without telephony']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3 User Characteristics', level=2)
add_table(['User Role', 'Description'], [
    ('Administrator', 'Full access to all features, manages users, campaigns, and system settings'),
    ('Agent', 'Views leads, monitors calls, reviews transcripts and analytics'),
    ('Developer', 'Interacts with API directly for testing and integration'),
])

doc.add_heading('2.4 Assumptions & Dependencies', level=2)
for a in ['PostgreSQL 16 is installed and running', 'Python 3.12+ and Node.js 18+ are available',
'Internet connection for Groq/Gemini/ElevenLabs APIs', 'ngrok or equivalent tunnel for SignalWire webhooks',
'FFmpeg available for audio processing (Whisper dependency)', 'SignalWire account with active phone number',
'At least 4GB RAM recommended for local Whisper model']:
    doc.add_paragraph(a, style='List Bullet')
doc.add_page_break()

# ── 3. System Features & Requirements ──
doc.add_heading('3. System Features & Requirements', level=1)

doc.add_heading('3.1 Authentication & Authorization', level=2)
doc.add_paragraph('REQ-AUTH-001: System shall authenticate users via email and password.')
doc.add_paragraph('REQ-AUTH-002: System shall issue JWT access tokens (60 min expiry) and refresh tokens (7 day expiry).')
doc.add_paragraph('REQ-AUTH-003: System shall automatically refresh expired tokens without user re-login.')
doc.add_paragraph('REQ-AUTH-004: System shall protect all API routes except /login and /register with JWT middleware.')
doc.add_paragraph('REQ-AUTH-005: Password shall be hashed using bcrypt before storage.')

doc.add_heading('3.2 Lead Management', level=2)
doc.add_paragraph('REQ-LEAD-001: System shall support creating, reading, updating, and deleting leads.')
doc.add_paragraph('REQ-LEAD-002: Each lead shall store name, phone, language, status, and custom data.')
doc.add_paragraph('REQ-LEAD-003: System shall support bulk lead upload via CSV file.')
doc.add_paragraph('REQ-LEAD-004: Leads shall be filterable by status, campaign, and search query.')
doc.add_paragraph('REQ-LEAD-005: System shall assign leads to campaigns and track call history per lead.')

doc.add_heading('3.3 Campaign Management', level=2)
doc.add_paragraph('REQ-CAMP-001: System shall support campaign creation with name, script template, greeting, and closing.')
doc.add_paragraph('REQ-CAMP-002: Campaigns shall have statuses: draft, active, paused, completed.')
doc.add_paragraph('REQ-CAMP-003: System shall track total leads and completed calls per campaign.')

doc.add_heading('3.4 Call Management', level=2)
doc.add_paragraph('REQ-CALL-001: System shall initiate outbound calls via SignalWire API using Twilio SDK.')
doc.add_paragraph('REQ-CALL-002: System shall generate TwiML for greeting, response playback, and hangup.')
doc.add_paragraph('REQ-CALL-003: System shall process speech input from calls and generate AI responses.')
doc.add_paragraph('REQ-CALL-004: System shall provide test-local endpoint that bypasses telephony for AI pipeline testing.')
doc.add_paragraph('REQ-CALL-005: System shall store full call transcripts in the database.')
doc.add_paragraph('REQ-CALL-006: System shall track call status (initiated, ringing, answered, completed).')

doc.add_heading('3.5 AI Conversation Engine', level=2)
doc.add_paragraph('REQ-AI-001: System shall support multiple LLM providers with automatic fallback on failure.')
doc.add_paragraph('REQ-AI-002: Providers shall be tried in order: Gemini, Groq, Ollama, OpenAI, Anthropic.')
doc.add_paragraph('REQ-AI-003: System shall use a warm, professional English speaking style for US health insurance leads.')
doc.add_paragraph('REQ-AI-004: System shall handle objection scenarios (not interested, price concern, busy).')
doc.add_paragraph('REQ-AI-005: System shall assess lead interest level (High/Medium/Low) after conversation.')
doc.add_paragraph('REQ-AI-006: System shall generate conversation summaries and sentiment scores.')

doc.add_heading('3.6 Text-to-Speech (TTS)', level=2)
doc.add_paragraph('REQ-TTS-001: System shall synthesize AI responses to spoken audio.')
doc.add_paragraph('REQ-TTS-002: System shall try ElevenLabs API first, fall back to gTTS on failure.')
doc.add_paragraph('REQ-TTS-003: Audio shall be served via HTTP endpoint for browser playback.')
doc.add_paragraph('REQ-TTS-004: Audio format shall be MP3 for browser compatibility.')

doc.add_heading('3.7 Speech-to-Text (STT)', level=2)
doc.add_paragraph('REQ-STT-001: System shall transcribe audio input to text using OpenAI Whisper (local).')
doc.add_paragraph('REQ-STT-002: System shall support English language transcription.')
doc.add_paragraph('REQ-STT-003: System shall support microphone recording for local test mode.')

doc.add_heading('3.8 Real-Time Dashboard', level=2)
doc.add_paragraph('REQ-DASH-001: Dashboard shall display total calls, leads, conversion rate, avg duration, avg sentiment.')
doc.add_paragraph('REQ-DASH-002: Dashboard shall include lead status breakdown chart (bar chart).')
doc.add_paragraph('REQ-DASH-003: Dashboard shall include daily call activity chart (line chart).')
doc.add_paragraph('REQ-DASH-004: Dashboard shall update in real-time via WebSocket when calls complete.')
doc.add_paragraph('REQ-DASH-005: Dashboard shall include a Test Call feature with audio playback.')

doc.add_heading('3.9 RAG Knowledge Base', level=2)
doc.add_paragraph('REQ-RAG-001: System shall index insurance documents for retrieval-augmented generation.')
doc.add_paragraph('REQ-RAG-002: System shall query knowledge base to enhance LLM responses with context.')
doc.add_paragraph('REQ-RAG-003: System shall support document addition and deletion.')

doc.add_heading('3.10 Analytics & Reporting', level=2)
doc.add_paragraph('REQ-ANALYTICS-001: System shall provide daily call statistics for the last 7/14/30/90 days.')
doc.add_paragraph('REQ-ANALYTICS-002: System shall calculate conversion rate from lead statuses.')
doc.add_paragraph('REQ-ANALYTICS-003: System shall compute average call duration and sentiment scores.')
doc.add_page_break()

# ── 4. External Interface Requirements ──
doc.add_heading('4. External Interface Requirements', level=1)
doc.add_heading('4.1 User Interfaces', level=2)
doc.add_paragraph('The system provides two frontend interfaces:')
doc.add_paragraph('React Web Application (Primary): Served on port 5173 during development. Includes Dashboard, Lead management, Campaign management, Call History, Settings, and Login pages.', style='List Bullet')
doc.add_paragraph('Streamlit Application (Secondary): Lightweight dashboard on port 8501 for quick monitoring.', style='List Bullet')

doc.add_heading('4.2 Hardware Interfaces', level=2)
doc.add_paragraph('REQ-HW-001: System requires microphone for local test mode STT recording.')
doc.add_paragraph('REQ-HW-002: System requires speakers/headphones for audio playback during testing.')
doc.add_paragraph('REQ-HW-003: Server requires minimum 2GB RAM, 4GB recommended for Whisper STT.')

doc.add_heading('4.3 Software Interfaces', level=2)
add_table(['External System', 'Interface Type', 'Protocol'], [
    ('PostgreSQL 16', 'Database', 'TCP 5432, asyncpg'),
    ('SignalWire', 'Telephony API', 'HTTPS REST + TwiML'),
    ('Groq API', 'LLM Inference', 'HTTPS REST, api.groq.com'),
    ('Gemini API', 'LLM Inference', 'HTTPS REST, generativelanguage.googleapis.com'),
    ('ElevenLabs API', 'TTS', 'HTTPS REST, api.elevenlabs.io'),
    ('Google Translate', 'gTTS', 'HTTPS, translate.google.com'),
])

doc.add_heading('4.4 Communication Interfaces', level=2)
doc.add_paragraph('REQ-COMM-001: Backend communicates with frontend via HTTP REST (JSON) on port 8000.')
doc.add_paragraph('REQ-COMM-002: Real-time updates use WebSocket protocol on /ws endpoint.')
doc.add_paragraph('REQ-COMM-003: SignalWire webhooks use HTTPS POST with TwiML XML responses.')
doc.add_paragraph('REQ-COMM-004: ngrok creates HTTPS tunnel from public URL to localhost:8000.')
doc.add_page_break()

# ── 5. Non-Functional Requirements ──
doc.add_heading('5. Non-Functional Requirements', level=1)
doc.add_heading('5.1 Performance', level=2)
doc.add_paragraph('REQ-PERF-001: LLM response time shall not exceed 30 seconds (including retries).')
doc.add_paragraph('REQ-PERF-002: TTS synthesis shall complete within 10 seconds for typical responses.')
doc.add_paragraph('REQ-PERF-003: API response time for non-LLM endpoints shall be under 500ms.')
doc.add_paragraph('REQ-PERF-004: System shall support up to 10 concurrent calls.')
doc.add_paragraph('REQ-PERF-005: Rate limiting shall allow 60 requests per minute per user.')
doc.add_heading('5.2 Security', level=2)
doc.add_paragraph('REQ-SEC-001: All API endpoints except login/register require JWT authentication.')
doc.add_paragraph('REQ-SEC-002: Passwords must be hashed with bcrypt before storage.')
doc.add_paragraph('REQ-SEC-003: JWT tokens expire after 60 minutes (access) and 7 days (refresh).')
doc.add_paragraph('REQ-SEC-004: CORS is configured to allow only specific origins (localhost, ngrok).')
doc.add_paragraph('REQ-SEC-005: API keys stored in .env file, never in code.')
doc.add_paragraph('REQ-SEC-006: Audio response endpoint does not require authentication (UUID-based).')
doc.add_heading('5.3 Reliability', level=2)
doc.add_paragraph('REQ-REL-001: LLM provider chain ensures service continues if one provider fails.')
doc.add_paragraph('REQ-REL-002: TTS falls back to gTTS if ElevenLabs is unavailable or quota exhausted.')
doc.add_paragraph('REQ-REL-003: Cache service degrades gracefully if Redis is unavailable.')
doc.add_paragraph('REQ-REL-004: Retry mechanism (3 attempts, exponential backoff) for LLM calls.')
doc.add_heading('5.4 Scalability', level=2)
doc.add_paragraph('REQ-SCALE-001: Backend is stateless and can scale horizontally behind a load balancer.')
doc.add_paragraph('REQ-SCALE-002: PostgreSQL handles concurrent connections via connection pooling.')
doc.add_paragraph('REQ-SCALE-003: WebSocket connections scale with server resources.')
doc.add_paragraph('REQ-SCALE-004: Audio storage is in-memory (not suitable for multi-worker scaling without Redis).')
doc.add_heading('5.5 Maintainability', level=2)
doc.add_paragraph('REQ-MAINT-001: Code follows consistent naming conventions (snake_case Python, camelCase TypeScript).')
doc.add_paragraph('REQ-MAINT-002: Services are decoupled and injectable (dependency injection pattern).')
doc.add_paragraph('REQ-MAINT-003: Configuration is centralized in .env file and config.py.')
doc.add_paragraph('REQ-MAINT-004: Database migrations use Alembic for schema versioning.')
doc.add_page_break()

# ── 6. System Architecture ──
doc.add_heading('6. System Architecture', level=1)
doc.add_paragraph('The system follows a three-tier architecture with a clear separation of concerns:')
add_table(['Tier', 'Port', 'Description'], [
    ('Presentation Tier (React Frontend)', '5173', 'Dashboard, Lead/Campaign management, Call History, Login. Real-time updates via WebSocket.'),
    ('Application Tier (FastAPI Backend)', '8000', 'REST API, WebSocket, JWT auth, LLM orchestration, TTS/STT, RAG engine. Stateless, scales horizontally.'),
    ('Data Tier (PostgreSQL 16)', '5432', 'Users, leads, campaigns, call logs, RAG documents. Async access via SQLAlchemy + asyncpg.'),
])
doc.add_paragraph()
doc.add_paragraph('External Integrations:')
for e in ['SignalWire - Outbound PSTN calls via Twilio SDK', 'Groq API - Primary LLM provider (Llama 3 70B, free tier)',
'Google Gemini API - Secondary LLM provider (free tier)', 'ElevenLabs API - Primary TTS (gTTS fallback)',
'Google Translate - gTTS backend for free TTS', 'ngrok - Public HTTPS tunnel for SignalWire webhooks']:
    doc.add_paragraph(e, style='List Bullet')
doc.add_page_break()

# ── 7. API Documentation ──
doc.add_heading('7. API Documentation', level=1)
doc.add_paragraph('Base URL: http://localhost:8000/api')
doc.add_heading('Authentication', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ('POST', '/auth/login', 'Login with email + password', 'None'),
    ('POST', '/auth/register', 'Register new user', 'None'),
    ('POST', '/auth/refresh', 'Refresh access token', 'None'),
])
doc.add_heading('Leads', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ('GET', '/leads', 'List leads (paginated, filterable)', 'JWT'),
    ('GET', '/leads/{id}', 'Get lead details', 'JWT'),
    ('POST', '/leads', 'Create lead', 'JWT'),
    ('PATCH', '/leads/{id}', 'Update lead', 'JWT'),
    ('DELETE', '/leads/{id}', 'Delete lead', 'JWT'),
])
doc.add_heading('Calls', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ('POST', '/calls/initiate', 'Start outbound call (SignalWire)', 'JWT'),
    ('POST', '/calls/test-local/{lead_id}', 'Start test call (no telephony)', 'JWT'),
    ('POST', '/calls/test-process-speech', 'Send text to AI pipeline', 'JWT'),
    ('GET', '/calls/audio-response/{id}', 'Get generated audio', 'None'),
    ('GET', '/calls/history', 'List call logs', 'JWT'),
    ('POST', '/calls/status', 'SignalWire status webhook', 'None'),
    ('POST', '/calls/outbound-twiml/{id}', 'TwiML greeting', 'None'),
])
doc.add_paragraph()
doc.add_paragraph('Full interactive API documentation available at http://localhost:8000/docs (Swagger UI).')
doc.add_page_break()

# ── 8. Database Design ──
doc.add_heading('8. Database Design', level=1)
doc.add_paragraph('The system uses PostgreSQL 16 with the following tables:')
doc.add_heading('Table: users', level=2)
add_table(['Column', 'Type', 'Constraints', 'Description'], [
    ('id', 'UUID', 'PK, default uuid4', 'User identifier'),
    ('email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'Login email'),
    ('password_hash', 'VARCHAR(255)', 'NOT NULL', 'bcrypt hash'),
    ('role', 'VARCHAR(50)', 'NOT NULL, default agent', 'User role'),
    ('created_at', 'TIMESTAMP', 'NOT NULL, default now()', 'Creation timestamp'),
])
doc.add_heading('Table: leads', level=2)
add_table(['Column', 'Type', 'Constraints', 'Description'], [
    ('id', 'UUID', 'PK, default uuid4', 'Lead identifier'),
    ('name', 'VARCHAR(255)', 'NOT NULL', 'Lead name'),
    ('phone', 'VARCHAR(50)', 'NOT NULL', 'Phone number'),
    ('language', 'VARCHAR(10)', 'NOT NULL', 'Language code'),
    ('status', 'VARCHAR(50)', 'NOT NULL, default pending', 'Lead status'),
    ('assigned_campaign_id', 'UUID', 'FK -> campaigns.id', 'Campaign assignment'),
    ('extra_data', 'JSONB', 'NULL', 'Custom metadata'),
    ('created_at', 'TIMESTAMP', 'NOT NULL', 'Creation timestamp'),
])
doc.add_heading('Table: call_logs', level=2)
add_table(['Column', 'Type', 'Constraints', 'Description'], [
    ('id', 'UUID', 'PK', 'Call log identifier'),
    ('lead_id', 'UUID', 'FK -> leads.id', 'Associated lead'),
    ('campaign_id', 'UUID', 'FK -> campaigns.id', 'Associated campaign'),
    ('direction', 'VARCHAR(10)', 'NOT NULL', 'outbound or inbound'),
    ('status', 'VARCHAR(50)', 'NOT NULL', 'Call status'),
    ('duration_seconds', 'INTEGER', 'NULL', 'Call duration'),
    ('transcript', 'JSONB', 'NULL', 'Conversation transcript'),
    ('summary', 'TEXT', 'NULL', 'AI-generated summary'),
    ('sentiment_score', 'FLOAT', 'NULL', '-1.0 to 1.0 sentiment'),
    ('created_at', 'TIMESTAMP', 'NOT NULL', 'Creation timestamp'),
])
doc.add_page_break()

# ── 9. Deployment ──
doc.add_heading('9. Deployment', level=1)
doc.add_heading('9.1 Local Development', level=2)
doc.add_paragraph('Backend: uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload')
doc.add_paragraph('Frontend: npm run dev (Vite, port 5173, proxies /api to backend)')
doc.add_paragraph('Database: PostgreSQL 16 on localhost:5432')
doc.add_paragraph('Tunnel: ngrok http 8000 for SignalWire webhooks')
doc.add_heading('9.2 Docker Compose', level=2)
doc.add_paragraph('docker-compose.yml provides PostgreSQL 16 service. Backend runs outside container for development.')
doc.add_paragraph('docker-compose.asterisk.yml provides Asterisk SIP server (requires Linux/WSL).')
doc.add_heading('9.3 Production (Render)', level=2)
doc.add_paragraph('render.yaml defines the deployment configuration.')
doc.add_paragraph('Backend service using Dockerfile, PostgreSQL 16 add-on, environment variables from .env.', style='List Bullet')
doc.add_paragraph('Health check at /health endpoint.', style='List Bullet')
doc.add_heading('9.4 Environment Variables Required', level=2)
doc.add_paragraph('Key variables: DATABASE_URL, GOOGLE_API_KEY, GROQ_API_KEY, SIGNALWIRE_SPACE, SIGNALWIRE_PROJECT_ID, SIGNALWIRE_AUTH_TOKEN, SIGNALWIRE_PHONE_NUMBER, SIGNALWIRE_WEBHOOK_BASE_URL, ELEVENLABS_API_KEY, JWT_SECRET_KEY')

# Save
output = os.path.join(os.path.dirname(__file__), 'SRS_AI_Health_Insurance_Cold_Calling_Agent.docx')
doc.save(output)
print(f'SRS saved to: {output}')
